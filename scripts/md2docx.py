"""Convert REQUIREMENTS.md to a formatted Word document."""
import re
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color):
    """Set cell background color."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_border(cell, **kwargs):
    """Set cell borders."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge, val in kwargs.items():
        element = OxmlElement(f'w:{edge}')
        element.set(qn('w:val'), val.get('val', 'single'))
        element.set(qn('w:sz'), val.get('sz', '4'))
        element.set(qn('w:color'), val.get('color', 'CCCCCC'))
        tcBorders.append(element)
    tcPr.append(tcBorders)

def add_hyperlink(paragraph, text, url):
    """Add a hyperlink to a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    c = OxmlElement('w:color')
    c.set(qn('w:val'), '0563C1')
    rPr.append(c)
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

# ── Style configuration ──────────────────────────────────────────
STYLES = {
    'heading1': {'size': 28, 'bold': True, 'color': '1A1A1A', 'space_before': 24, 'space_after': 12},
    'heading2': {'size': 22, 'bold': True, 'color': '2C3E50', 'space_before': 20, 'space_after': 10},
    'heading3': {'size': 18, 'bold': True, 'color': '34495E', 'space_before': 16, 'space_after': 8},
    'heading4': {'size': 15, 'bold': True, 'color': '555555', 'space_before': 14, 'space_after': 6},
    'body': {'size': 11, 'color': '333333', 'space_after': 6},
    'code': {'size': 9.5, 'color': '2D2D2D', 'font_name': 'Consolas'},
    'inline_code': {'size': 9.5, 'color': 'C0392B', 'font_name': 'Consolas'},
}

def apply_style(run, style_name):
    s = STYLES[style_name]
    run.font.size = Pt(s['size'])
    run.font.color.rgb = RGBColor.from_string(s['color'])
    if s.get('bold'):
        run.bold = True
    if 'font_name' in s:
        run.font.name = s['font_name']

def add_styled_paragraph(doc, text, style_name, alignment=None):
    p = doc.add_paragraph()
    if STYLES[style_name].get('space_before'):
        p.paragraph_format.space_before = Pt(STYLES[style_name]['space_before'])
    if STYLES[style_name].get('space_after'):
        p.paragraph_format.space_after = Pt(STYLES[style_name]['space_after'])
    if alignment:
        p.alignment = alignment
    run = p.add_run(text)
    apply_style(run, style_name)
    return p

def add_mixed_paragraph(doc, segments):
    """segments: list of (text, style_name_or_None) tuples."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    for text, style in segments:
        if text:
            run = p.add_run(text)
            if style:
                apply_style(run, style)
            else:
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor.from_string('333333')
    return p

def parse_inline(text):
    """Parse inline markdown: bold, italic, inline code, and plain text.
    Returns list of (text, style_or_None)."""
    segments = []
    # Patterns in order: inline code, bold, italic
    pattern = r'(`[^`]+`|\*\*[^*]+\*\*|\*[^*]+\*|\[[^\]]+\]\([^)]+\))'
    parts = re.split(pattern, text)
    for part in parts:
        if not part:
            continue
        if part.startswith('`') and part.endswith('`'):
            segments.append((part[1:-1], 'inline_code'))
        elif part.startswith('**') and part.endswith('**'):
            # bold
            inner = part[2:-2]
            # check if there's italic inside bold
            if '*' in inner:
                sub_parts = re.split(r'(\*[^*]+\*)', inner)
                for sp in sub_parts:
                    if sp.startswith('*') and sp.endswith('*'):
                        segments.append((sp[1:-1], 'bold_italic'))
                    else:
                        segments.append((sp, 'bold'))
            else:
                segments.append((inner, 'bold'))
        elif part.startswith('*') and part.endswith('*') and len(part) > 2:
            segments.append((part[1:-1], 'italic'))
        elif part.startswith('[') and '](' in part:
            m = re.match(r'\[([^\]]+)\]\(([^)]+)\)', part)
            if m:
                segments.append((m.group(1), 'link'))
        else:
            segments.append((part, None))
    return segments

def add_inline_paragraph(doc, text, space_after=6):
    """Add paragraph with inline markdown parsing."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    segments = parse_inline(text)
    for seg_text, style in segments:
        if not seg_text:
            continue
        run = p.add_run(seg_text)
        if style == 'inline_code':
            apply_style(run, 'inline_code')
        elif style == 'bold':
            run.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor.from_string('333333')
        elif style == 'bold_italic':
            run.bold = True
            run.italic = True
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor.from_string('333333')
        elif style == 'italic':
            run.italic = True
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor.from_string('333333')
        elif style == 'link':
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor.from_string('0563C1')
            run.underline = True
        else:
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor.from_string('333333')
    return p

def add_code_block(doc, code_text):
    """Add a code block with background shading."""
    for line in code_text.strip().split('\n'):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Cm(0.5)
        # Add shading to paragraph
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:fill'), 'F5F5F5')
        pPr.append(shd)
        run = p.add_run(line if line else ' ')
        apply_style(run, 'code')

def add_table_from_md(doc, lines, start_idx):
    """Parse a markdown table and add it to the document. Returns end index."""
    # Find all table rows
    table_lines = []
    i = start_idx
    while i < len(lines) and lines[i].strip().startswith('|'):
        table_lines.append(lines[i].strip())
        i += 1

    if len(table_lines) < 2:
        return start_idx + 1

    # Parse header
    headers = [c.strip() for c in table_lines[0].split('|')[1:-1]]
    # Skip separator line (index 1)
    rows = []
    for tl in table_lines[2:]:
        rows.append([c.strip() for c in tl.split('|')[1:-1]])

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr_cells = table.rows[0].cells
    for j, h in enumerate(headers):
        hdr_cells[j].text = ''
        p = hdr_cells[j].paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor.from_string('FFFFFF')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(hdr_cells[j], '2C3E50')

    # Data rows
    for r_idx, row in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx, cell_text in enumerate(row):
            cells[c_idx].text = ''
            p = cells[c_idx].paragraphs[0]
            # Parse inline in cell
            segments = parse_inline(cell_text)
            for seg_text, style in segments:
                run = p.add_run(seg_text)
                run.font.size = Pt(10)
                if style in ('inline_code',):
                    apply_style(run, 'inline_code')
                elif style == 'bold':
                    run.bold = True
                elif style == 'italic':
                    run.italic = True
            if r_idx % 2 == 0:
                set_cell_shading(cells[c_idx], 'F9F9F9')

    doc.add_paragraph()  # space after table
    return i

# ── Main conversion ──────────────────────────────────────────────

def convert_md_to_docx(md_path, docx_path):
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    style.font.size = Pt(11)
    style.font.name = 'Georgia'
    style.font.color.rgb = RGBColor.from_string('333333')

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    lines = content.split('\n')
    i = 0
    in_code_block = False
    code_buffer = []
    in_table = False

    while i < len(lines):
        line = lines[i]

        # Code block fence
        if line.strip().startswith('```'):
            if in_code_block:
                add_code_block(doc, '\n'.join(code_buffer))
                code_buffer = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_buffer.append(line)
            i += 1
            continue

        # Horizontal rule
        if line.strip() == '---':
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(12)
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '6')
            bottom.set(qn('w:color'), 'CCCCCC')
            pBdr.append(bottom)
            pPr.append(pBdr)
            i += 1
            continue

        # Table
        if line.strip().startswith('|') and not in_code_block:
            # Check if it's a real table (has separator row next)
            if i + 1 < len(lines) and lines[i + 1].strip().startswith('|') and '---' in lines[i + 1]:
                i = add_table_from_md(doc, lines, i)
                continue

        # Headings
        h1 = re.match(r'^# (.+)$', line)
        h2 = re.match(r'^## (.+)$', line)
        h3 = re.match(r'^### (.+)$', line)
        h4 = re.match(r'^#### (.+)$', line)

        if h1:
            text = h1.group(1)
            # Remove anchor links like {#id}
            text = re.sub(r'\{#[^}]+\}', '', text).strip()
            add_styled_paragraph(doc, text, 'heading1')
            i += 1
            continue
        if h2:
            text = h2.group(1)
            text = re.sub(r'\{#[^}]+\}', '', text).strip()
            add_styled_paragraph(doc, text, 'heading2')
            i += 1
            continue
        if h3:
            text = h3.group(1)
            text = re.sub(r'\{#[^}]+\}', '', text).strip()
            add_styled_paragraph(doc, text, 'heading3')
            i += 1
            continue
        if h4:
            text = h4.group(1)
            text = re.sub(r'\{#[^}]+\}', '', text).strip()
            add_styled_paragraph(doc, text, 'heading4')
            i += 1
            continue

        # Bullet list
        bullet_match = re.match(r'^(\s*)[-*]\s+(.+)$', line)
        if bullet_match:
            indent_level = len(bullet_match.group(1))
            text = bullet_match.group(2)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.0 + indent_level * 0.5)
            p.paragraph_format.space_after = Pt(3)
            segments = parse_inline(text)
            for seg_text, style in segments:
                run = p.add_run(seg_text)
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor.from_string('333333')
                if style == 'inline_code':
                    apply_style(run, 'inline_code')
                elif style == 'bold':
                    run.bold = True
                elif style == 'italic':
                    run.italic = True
                elif style == 'bold_italic':
                    run.bold = True
                    run.italic = True
            i += 1
            continue

        # Numbered list (simple — just indent)
        numbered_match = re.match(r'^(\s*)\d+\.\s+(.+)$', line)
        if numbered_match:
            indent_level = len(numbered_match.group(1))
            text = numbered_match.group(2)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.0 + indent_level * 0.5)
            p.paragraph_format.space_after = Pt(3)
            segments = parse_inline(text)
            for seg_text, style in segments:
                run = p.add_run(seg_text)
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor.from_string('333333')
                if style == 'inline_code':
                    apply_style(run, 'inline_code')
                elif style == 'bold':
                    run.bold = True
                elif style == 'italic':
                    run.italic = True
                elif style == 'bold_italic':
                    run.bold = True
                    run.italic = True
            i += 1
            continue

        # Empty line
        if not line.strip():
            i += 1
            continue

        # Regular paragraph with inline formatting
        add_inline_paragraph(doc, line)
        i += 1

    # Handle any remaining code buffer
    if code_buffer:
        add_code_block(doc, '\n'.join(code_buffer))

    doc.save(docx_path)
    print(f"Word document saved to: {docx_path}")


if __name__ == '__main__':
    import sys
    md_file = sys.argv[1] if len(sys.argv) > 1 else 'REQUIREMENTS.md'
    docx_file = sys.argv[2] if len(sys.argv) > 2 else md_file.replace('.md', '.docx')
    convert_md_to_docx(md_file, docx_file)
